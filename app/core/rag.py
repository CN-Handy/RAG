import logging
import datetime
from typing import List, Dict, Any

import numpy as np
import pdfplumber
import torch
from openai import OpenAI
from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer, AutoModelForSequenceClassification

from app.config import config, resolve_path
from app.db.vector_store import es

logger = logging.getLogger(__name__)

device = config["device"]

EMBEDDING_MODEL_PARAMS: Dict[str, Any] = {}

BASIC_QA_TEMPLATE = (
    "现在的时间是{time}。你是一个专家，擅长回答用户提问。"
    "请结合下面给定的资料回答问题。\n"
    "如果问题无法从资料中获得答案，请回答【无法回答】。\n\n"
    "资料：\n{context}\n\n"
    "问题：{question}\n"
)

QUERY_REWRITE_TEMPLATE = (
    "请将以下用户问题改写为适合文档检索的关键词列表，关键词之间用空格分隔。"
    "只输出关键词，不要输出其他内容。\n\n"
    "用户问题：{query}\n\n"
    "关键词列表："
)


def _load_embedding_model(model_name: str, model_path: str) -> None:
    global EMBEDDING_MODEL_PARAMS
    if model_name in ["bge-small-zh-v1.5", "bge-base-zh-v1.5"]:
        logger.info(f"加载 embedding 模型: {model_path}")
        EMBEDDING_MODEL_PARAMS["embedding_model"] = SentenceTransformer(model_path)
        logger.info("Embedding 模型加载完成")


def _load_rerank_model(model_name: str, model_path: str) -> None:
    global EMBEDDING_MODEL_PARAMS
    if model_name in ["bge-reranker-base"]:
        logger.info(f"加载 rerank 模型: {model_path}")
        EMBEDDING_MODEL_PARAMS["rerank_model"] = AutoModelForSequenceClassification.from_pretrained(model_path)
        EMBEDDING_MODEL_PARAMS["rerank_tokenizer"] = AutoTokenizer.from_pretrained(model_path)
        EMBEDDING_MODEL_PARAMS["rerank_model"].eval()
        EMBEDDING_MODEL_PARAMS["rerank_model"].to(device)
        logger.info("Rerank 模型加载完成")


if config["rag"]["use_embedding"]:
    _name = config["rag"]["embedding_model"]
    _path = resolve_path(config["models"]["embedding_model"][_name]["local_url"])
    try:
        _load_embedding_model(_name, _path)
    except Exception as _e:
        logger.warning(f"Embedding 模型加载失败（{_path}）: {_e}。请将模型文件放入 models/BAAI/ 目录。")

if config["rag"]["use_rerank"]:
    _name = config["rag"]["rerank_model"]
    _path = resolve_path(config["models"]["rerank_model"][_name]["local_url"])
    try:
        _load_rerank_model(_name, _path)
    except Exception as _e:
        logger.warning(f"Rerank 模型加载失败（{_path}）: {_e}。请将模型文件放入 models/BAAI/ 目录。")


def split_text_with_overlap(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        chunks.append(text[start: start + chunk_size])
        start += chunk_size - chunk_overlap
    return chunks


class RAG:
    def __init__(self):
        self.embedding_model = config["rag"]["embedding_model"]
        self.rerank_model = config["rag"]["rerank_model"]
        self.use_rerank = config["rag"]["use_rerank"]
        self.embedding_dims = config["models"]["embedding_model"][self.embedding_model]["dims"]
        self.chunk_size = config["rag"]["chunk_size"]
        self.chunk_overlap = config["rag"]["chunk_overlap"]
        self.chunk_candidate = config["rag"]["chunk_candidate"]
        self.client = OpenAI(
            api_key=config["rag"]["llm_api_key"],
            base_url=config["rag"]["llm_base"],
        )
        self.llm_model = config["rag"]["llm_model"]

    # ------------------------------------------------------------------
    # 文档提取
    # ------------------------------------------------------------------

    def _extract_pdf_content(self, knowledge_id: int, document_id: int, title: str, file_path: str) -> bool:
        try:
            pdf = pdfplumber.open(file_path)
        except Exception as e:
            logger.error(f"打开 PDF 失败: {file_path} — {e}")
            return False

        logger.info(f"PDF 提取开始: {file_path}，共 {len(pdf.pages)} 页")
        abstract = ""

        for page_number, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            if page_number < 3:
                abstract += "\n" + page_text

            # 整页向量（chunk_id=0）
            page_vec = self.get_embedding(page_text)
            es.index(index="chunk_info", document={
                "document_id": document_id, "knowledge_id": knowledge_id,
                "page_number": page_number, "chunk_id": 0,
                "chunk_content": page_text, "chunk_images": [], "chunk_tables": [],
                "embedding_vector": page_vec,
            })

            # 分块向量
            chunks = split_text_with_overlap(page_text, self.chunk_size, self.chunk_overlap)
            chunk_vecs = self.get_embedding(chunks)
            for idx, chunk_text in enumerate(chunks):
                es.index(index="chunk_info", document={
                    "document_id": document_id, "knowledge_id": knowledge_id,
                    "page_number": page_number, "chunk_id": idx + 1,
                    "chunk_content": chunk_text, "chunk_images": [], "chunk_tables": [],
                    "embedding_vector": chunk_vecs[idx],
                })

        es.index(index="document_meta", document={
            "document_id": document_id, "knowledge_id": knowledge_id,
            "document_name": title, "file_path": file_path, "abstract": abstract,
        })
        logger.info(f"PDF 提取完成: {file_path}")
        return True

    def _extract_word_content(self, knowledge_id: int, document_id: int, title: str, file_path: str) -> bool:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
        except Exception as e:
            logger.error(f"打开 Word 文件失败: {file_path} — {e}")
            return False

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            logger.warning(f"Word 文档无文本内容: {file_path}")
            return False

        abstract = "\n".join(paragraphs[:3])
        for page_number, para_text in enumerate(paragraphs):
            para_vec = self.get_embedding(para_text)
            es.index(index="chunk_info", document={
                "document_id": document_id, "knowledge_id": knowledge_id,
                "page_number": page_number, "chunk_id": 0,
                "chunk_content": para_text, "chunk_images": [], "chunk_tables": [],
                "embedding_vector": para_vec,
            })
            chunks = split_text_with_overlap(para_text, self.chunk_size, self.chunk_overlap)
            if len(chunks) > 1:
                chunk_vecs = self.get_embedding(chunks)
                for idx, chunk_text in enumerate(chunks):
                    es.index(index="chunk_info", document={
                        "document_id": document_id, "knowledge_id": knowledge_id,
                        "page_number": page_number, "chunk_id": idx + 1,
                        "chunk_content": chunk_text, "chunk_images": [], "chunk_tables": [],
                        "embedding_vector": chunk_vecs[idx],
                    })

        es.index(index="document_meta", document={
            "document_id": document_id, "knowledge_id": knowledge_id,
            "document_name": title, "file_path": file_path, "abstract": abstract,
        })
        logger.info(f"Word 提取完成: {file_path}，共 {len(paragraphs)} 段")
        return True

    def _update_parse_status(self, document_id: int, status: str) -> None:
        from app.db.models import KnowledgeDocument, Session
        try:
            with Session() as session:
                doc = session.query(KnowledgeDocument).filter(
                    KnowledgeDocument.document_id == document_id
                ).first()
                if doc:
                    doc.parse_status = status
                    session.commit()
        except Exception as e:
            logger.warning(f"更新 parse_status 失败 document_id={document_id}: {e}")

    def extract_content(self, knowledge_id: int, document_id: int, title: str, file_type: str, file_path: str) -> None:
        success = False
        try:
            if "pdf" in file_type:
                success = self._extract_pdf_content(knowledge_id, document_id, title, file_path)
            elif "word" in file_type or "docx" in file_type or "vnd.openxmlformats" in file_type:
                success = self._extract_word_content(knowledge_id, document_id, title, file_path)
            else:
                logger.warning(f"不支持的文件类型: {file_type}，跳过提取")
        except Exception as e:
            logger.error(f"文档提取异常 document_id={document_id}: {e}", exc_info=True)

        final_status = "completed" if success else "failed"
        self._update_parse_status(document_id, final_status)
        logger.info(f"提取任务结束 document_id={document_id} parse_status={final_status}")

    # ------------------------------------------------------------------
    # 模型推理
    # ------------------------------------------------------------------

    def get_embedding(self, text) -> np.ndarray:
        if self.embedding_model in ["bge-small-zh-v1.5", "bge-base-zh-v1.5"]:
            return EMBEDDING_MODEL_PARAMS["embedding_model"].encode(text, normalize_embeddings=True)
        raise NotImplementedError(f"不支持的 embedding 模型: {self.embedding_model}")

    def get_rank(self, text_pair: List[List[str]]) -> np.ndarray:
        if self.rerank_model in ["bge-reranker-base"]:
            with torch.no_grad():
                inputs = EMBEDDING_MODEL_PARAMS["rerank_tokenizer"](
                    text_pair, padding=True, truncation=True,
                    return_tensors="pt", max_length=512,
                )
                inputs = {k: v.to(device) for k, v in inputs.items()}
                scores = EMBEDDING_MODEL_PARAMS["rerank_model"](**inputs, return_dict=True).logits.view(-1).float()
                return scores.cpu().numpy()
        raise NotImplementedError(f"不支持的 rerank 模型: {self.rerank_model}")

    # ------------------------------------------------------------------
    # 检索
    # ------------------------------------------------------------------

    def query_document(self, query: str, knowledge_id: int) -> List[Dict]:
        # BM25 全文检索
        word_resp = es.search(
            index="chunk_info",
            body={
                "query": {
                    "bool": {
                        "must": [{"match": {"chunk_content": query}}],
                        "filter": [{"term": {"knowledge_id": knowledge_id}}],
                    }
                },
                "size": 50,
            },
            fields=["chunk_id", "document_id", "knowledge_id", "page_number", "chunk_content"],
            source=False,
        )

        # KNN 语义检索
        query_vec = self.get_embedding(query)
        knn_resp = es.search(
            index="chunk_info",
            knn={
                "field": "embedding_vector",
                "query_vector": query_vec,
                "k": 50,
                "num_candidates": 100,
                "filter": {"term": {"knowledge_id": knowledge_id}},
            },
            fields=["chunk_id", "document_id", "knowledge_id", "page_number", "chunk_content"],
            source=False,
        )

        # RRF 融合（k=60 来自原始论文）
        k = 60
        fusion_score: Dict[str, float] = {}
        id2record: Dict[str, Dict] = {}

        for idx, hit in enumerate(word_resp["hits"]["hits"]):
            _id = hit["_id"]
            fusion_score[_id] = fusion_score.get(_id, 0) + 1 / (idx + k)
            id2record.setdefault(_id, hit["fields"])

        for idx, hit in enumerate(knn_resp["hits"]["hits"]):
            _id = hit["_id"]
            fusion_score[_id] = fusion_score.get(_id, 0) + 1 / (idx + k)
            id2record.setdefault(_id, hit["fields"])

        sorted_ids = sorted(fusion_score, key=lambda x: fusion_score[x], reverse=True)
        sorted_records = [id2record[i] for i in sorted_ids][: self.chunk_candidate]
        sorted_content = [r["chunk_content"] for r in sorted_records]

        # 可选 rerank
        if self.use_rerank:
            text_pair = [[query, c[0]] for c in sorted_content]   # ES fields 返回 list，取 [0]
            rerank_score = self.get_rank(text_pair)
            rerank_idx = np.argsort(rerank_score)[::-1]
            sorted_records = [sorted_records[i] for i in rerank_idx]

        return sorted_records

    # ------------------------------------------------------------------
    # 查询改写（提升检索召回率，面试亮点）
    # ------------------------------------------------------------------

    def query_rewrite(self, query: str, history: List[Dict] = None) -> str:
        """用 LLM 将问题改写为检索关键词，多轮时结合历史上下文解析指代。"""
        try:
            if history:
                history_text = "\n".join(
                    f"用户：{m['content']}" for m in history if m["role"] == "user"
                )
                prompt = (
                    f"对话历史：\n{history_text}\n\n"
                    f"当前问题：{query}\n\n"
                    "请结合对话历史，将当前问题改写为独立的、适合文档检索的关键词列表，"
                    "关键词之间用空格分隔，只输出关键词：\n\n关键词列表："
                )
            else:
                prompt = QUERY_REWRITE_TEMPLATE.format(query=query)

            rewritten = self.chat(
                [{"role": "user", "content": prompt}], top_p=0.9, temperature=0.1
            ).content.strip()
            logger.info(f"查询改写: '{query}' -> '{rewritten}'")
            return rewritten if rewritten else query
        except Exception as e:
            logger.warning(f"查询改写失败，使用原始查询: {e}")
            return query

    def query_parse(self, query: str) -> str:
        """解析查询意图和关键实体（扩展接口，暂返回原始查询）。"""
        return query

    # ------------------------------------------------------------------
    # 对话
    # ------------------------------------------------------------------

    def chat_with_rag(self, knowledge_id: int, messages: List[Dict]) -> List[Dict]:
        query = messages[-1]["content"]
        # 有历史轮次时传入，用于解析"它"、"上面提到的"等指代
        history = messages[:-1] if len(messages) > 1 else None

        retrieval_query = self.query_rewrite(query, history=history)
        related_records = self.query_document(retrieval_query, knowledge_id)
        logger.info(f"检索到 {len(related_records)} 条相关片段")

        context = "\n".join(r["chunk_content"][0] for r in related_records)
        rag_prompt = BASIC_QA_TEMPLATE.format(
            time=str(datetime.datetime.now()),
            context=context,
            question=query,   # 原始问题展示给 LLM
        )

        # 将 RAG prompt 替换最后一条用户消息，历史保持原样
        llm_messages = (history or []) + [{"role": "user", "content": rag_prompt}]
        response = self.chat(llm_messages, top_p=0.7, temperature=0.9).content
        messages.append({"role": "assistant", "content": response})

        return messages

    def chat(self, messages: List[Dict], top_p: float, temperature: float) -> Any:
        completion = self.client.chat.completions.create(
            model=self.llm_model,
            messages=messages,
            top_p=top_p,
            temperature=temperature,
        )
        return completion.choices[0].message
